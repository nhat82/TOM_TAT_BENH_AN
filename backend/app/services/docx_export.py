"""
Build a formatted DOCX from a Vietnamese medical summary string.

Section detection: lines matching "^\\d+\\.\\s" or "^\\*\\*\\d+\\.\\s" are
treated as numbered headings; everything else becomes body paragraph text.
"""

from __future__ import annotations

import io
import logging
import re
from datetime import date, datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlalchemy import text

from app.services.database import db

log = logging.getLogger(__name__)

# ── patient info helpers ──────────────────────────────────────────────────────

_NULL_VALUES = {"nan", "none", "", "0", "0.0", "0001-01-01 00:00:00"}


def _clean(val) -> str:
    if val is None:
        return ""
    s = str(val).strip()
    return "" if s.lower() in _NULL_VALUES else s


def _fmt_date(val) -> str:
    """Return dd/mm/yyyy string from a date/datetime/string value, or ''."""
    if val is None:
        return ""
    s = str(val).strip()
    if s.lower() in _NULL_VALUES:
        return ""
    try:
        return datetime.strptime(s[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
    except ValueError:
        return s


def fetch_patient_info(pid: str) -> dict:
    """Query medical_records for patient demographics; returns display-ready strings."""
    try:
        with db._engine.connect() as conn:
            result = conn.execute(
                text(
                    "SELECT birthdayyear, "
                    "dm_tinhcode, medicalrecorddate_in, medicalrecorddate_out "
                    "FROM medical_records WHERE ma_bn_an = :pid LIMIT 1"
                ),
                {"pid": pid},
            )
            row = result.mappings().first()
    except Exception:
        log.warning("Could not fetch patient demographics for %s", pid)
        return {}

    if not row:
        return {}

    gender_map = {"1": "Nam", "2": "Nữ"}
    birth_year = _clean(row.get("birthdayyear"))
    age = ""
    if birth_year:
        try:
            age = str(date.today().year - int(birth_year))
        except ValueError:
            pass

    gender_raw = _clean(row.get("dm_gioitinhid"))
    gender = gender_map.get(gender_raw, gender_raw)

    _na = "N/A"
    return {
        "patient_name":   _clean(row.get("ho_ten"))                    or _na,
        "birthday":       birth_year                                    or _na,
        "age":            age                                           or _na,
        "gender":         gender                                        or _na,
        "id_number":      _clean(row.get("cccd"))                      or _na,
        "province":        _clean(row.get("dm_tinhcode"))                or _na,
        "admission_date": _fmt_date(row.get("medicalrecorddate_in"))    or _na,
        "discharge_date": _fmt_date(row.get("medicalrecorddate_out"))   or _na,
    }


_SECTION_RE = re.compile(r"^\*{0,2}(\d+)\.\s+(.+?)\*{0,2}$")

# ── helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BDBDBD")
    pb.append(bottom)
    pPr.append(pb)


# ── parser ────────────────────────────────────────────────────────────────────

def _parse_summary(text: str) -> list[tuple[str, str]]:
    """
    Return a list of (kind, content) tuples:
      ("heading", "1. Section title")
      ("body",    "Paragraph text …")
    """
    blocks: list[tuple[str, str]] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        m = _SECTION_RE.match(line)
        if m:
            blocks.append(("heading", f"{m.group(1)}. {m.group(2)}"))
        else:
            blocks.append(("body", line))
    return blocks


# ── public API ────────────────────────────────────────────────────────────────

_PATIENT_LABELS = [
    ("patient_name",   "Họ tên"),
    ("birthday",       "Năm sinh"),
    ("age",            "Tuổi"),
    ("gender",         "Giới tính"),
    ("ethnicity",      "Dân tộc"),
    ("address",        "Địa chỉ"),
    ("id_number",      "Số CMND/CCCD"),
    ("admission_date", "Ngày nhập viện"),
    ("discharge_date", "Ngày xuất viện"),
]


def _add_patient_info_table(doc: Document, info: dict) -> None:
    rows = [(label, info.get(key, "")) for key, label in _PATIENT_LABELS if info.get(key)]
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    col_widths = (Inches(1.8), Inches(4.0))

    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]

        _set_cell_bg(row.cells[0], "EAF2FF")

        label_p = row.cells[0].paragraphs[0]
        label_run = label_p.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(10)

        value_p = row.cells[1].paragraphs[0]
        value_run = value_p.add_run(value)
        value_run.font.size = Pt(10)

    doc.add_paragraph()  # spacer after table


def build_docx(patient_id: str, summary: str, patient_info: dict | None = None) -> bytes:
    """
    Render *summary* as a formatted DOCX and return raw bytes.
    *patient_info* keys: patient_name, birthday, age, gender, ethnicity,
                         address, id_number, admission_date, discharge_date
    """
    doc = Document()

    # ── page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── title block ───────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("TÓM TẮT BỆNH ÁN")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(2)
    mr = meta.add_run(f"Mã bệnh nhân: {patient_id}    |    Ngày xuất: {date.today().strftime('%d/%m/%Y')}")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    _add_horizontal_rule(doc)
    doc.add_paragraph()  # spacer

    # ── patient demographics table ────────────────────────────────────────────
    if patient_info:
        _add_patient_info_table(doc, patient_info)

    # ── body ──────────────────────────────────────────────────────────────────
    for kind, content in _parse_summary(summary):
        if kind == "heading":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        else:
            p = doc.add_paragraph(content)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(4)
            for run in p.runs:
                run.font.size = Pt(11)

    # ── footer ────────────────────────────────────────────────────────────────
    _add_horizontal_rule(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("Tài liệu được tạo tự động")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)
    fr.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
